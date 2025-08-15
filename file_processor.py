import io
import os
from typing import Union
import streamlit as st

# Import libraries for file processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

class FileProcessor:
    def __init__(self):
        self.supported_types = ['txt', 'pdf', 'docx', 'csv']
    
    def process_file(self, uploaded_file) -> str:
        """
        Process uploaded file and extract text content
        """
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            if file_type == 'txt':
                return self._process_txt(uploaded_file)
            elif file_type == 'pdf':
                return self._process_pdf(uploaded_file)
            elif file_type == 'docx':
                return self._process_docx(uploaded_file)
            elif file_type == 'csv':
                return self._process_csv(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error processing {file_type.upper()} file: {str(e)}")
    
    def _process_txt(self, uploaded_file) -> str:
        """
        Process text files
        """
        try:
            # Try different encodings
            content = uploaded_file.read()
            
            # Try UTF-8 first
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                # Try latin-1 as fallback
                try:
                    return content.decode('latin-1')
                except UnicodeDecodeError:
                    # Try with error handling
                    return content.decode('utf-8', errors='replace')
        except Exception as e:
            raise Exception(f"Could not read text file: {str(e)}")
    
    def _process_pdf(self, uploaded_file) -> str:
        """
        Process PDF files using PyPDF2
        """
        if PyPDF2 is None:
            raise Exception("PyPDF2 is required for PDF processing. Please install it.")
        
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the PDF. The PDF might contain only images or be password protected.")
            
            return extracted_text
            
        except Exception as e:
            if "No text could be extracted" in str(e):
                raise e
            else:
                raise Exception(f"Error reading PDF file: {str(e)}")
    
    def _process_docx(self, uploaded_file) -> str:
        """
        Process DOCX files using python-docx
        """
        if Document is None:
            raise Exception("python-docx is required for DOCX processing. Please install it.")
        
        try:
            doc = Document(uploaded_file)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n\n'.join(text_content)
            
            if not extracted_text.strip():
                raise Exception("No text could be extracted from the DOCX file.")
            
            return extracted_text
            
        except Exception as e:
            if "No text could be extracted" in str(e):
                raise e
            else:
                raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _process_csv(self, uploaded_file) -> str:
        """
        Process CSV files using pandas
        """
        if pd is None:
            raise Exception("pandas is required for CSV processing. Please install it.")
        
        try:
            # Try to read CSV with different encodings
            try:
                df = pd.read_csv(uploaded_file, encoding='utf-8')
            except UnicodeDecodeError:
                uploaded_file.seek(0)  # Reset file pointer
                df = pd.read_csv(uploaded_file, encoding='latin-1')
            
            if df.empty:
                raise Exception("The CSV file is empty.")
            
            # Convert DataFrame to text format suitable for quiz generation
            text_content = []
            
            # Add column headers as context
            text_content.append("Data Table Information:")
            text_content.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_content.append(f"Total Records: {len(df)}")
            text_content.append("")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_content.append("Numeric Data Summary:")
                for col in numeric_cols:
                    stats = df[col].describe()
                    text_content.append(f"{col}: Mean={stats['mean']:.2f}, Min={stats['min']:.2f}, Max={stats['max']:.2f}")
                text_content.append("")
            
            # Add categorical data summary
            categorical_cols = df.select_dtypes(include=['object']).columns
            if len(categorical_cols) > 0:
                text_content.append("Categorical Data Summary:")
                for col in categorical_cols:
                    unique_vals = df[col].nunique()
                    top_values = df[col].value_counts().head(3)
                    text_content.append(f"{col}: {unique_vals} unique values. Top values: {', '.join([f'{k}({v})' for k, v in top_values.items()])}")
                text_content.append("")
            
            # Add sample rows for context (first 5 rows)
            text_content.append("Sample Data (First 5 Rows):")
            for idx, row in df.head(5).iterrows():
                row_text = []
                for col in df.columns:
                    row_text.append(f"{col}: {row[col]}")
                text_content.append(f"Row {int(idx) + 1}: " + " | ".join(row_text))
            
            # Add data patterns and insights if possible
            text_content.append("")
            text_content.append("Data Insights:")
            text_content.append(f"This dataset contains {len(df)} records with {len(df.columns)} columns.")
            
            if len(numeric_cols) > 0:
                text_content.append(f"Numeric columns ({len(numeric_cols)}): {', '.join(numeric_cols)}")
            
            if len(categorical_cols) > 0:
                text_content.append(f"Categorical columns ({len(categorical_cols)}): {', '.join(categorical_cols)}")
            
            extracted_text = '\n'.join(text_content)
            
            if not extracted_text.strip():
                raise Exception("No meaningful data could be extracted from the CSV file.")
            
            return extracted_text
            
        except Exception as e:
            if "No meaningful data could be extracted" in str(e) or "empty" in str(e):
                raise e
            else:
                raise Exception(f"Error reading CSV file: {str(e)}")
    
    def validate_content_length(self, content: str, min_length: int = 50) -> bool:
        """
        Validate that content has sufficient length for quiz generation
        """
        return len(content.strip()) >= min_length
    
    def get_content_stats(self, content: str) -> dict:
        """
        Get basic statistics about the content
        """
        words = content.split()
        sentences = content.split('.')
        
        return {
            'characters': len(content),
            'words': len(words),
            'sentences': len([s for s in sentences if s.strip()]),
            'paragraphs': len([p for p in content.split('\n\n') if p.strip()])
        }

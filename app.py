import streamlit as st
import os
import json
import io
import re
import random

# Import libraries with error handling
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

# Streamlit page config
st.set_page_config(
    page_title="AI Quiz Generator",
    page_icon="📚",
    layout="wide"
)

# Platform compatibility check
def check_platform_compatibility():
    """Check if we're running on a compatible platform"""
    import platform
    system = platform.system().lower()
    
    # Show platform info in sidebar
    with st.sidebar:
        st.info(f"🖥️ Running on: {platform.system()}")
        
        # API Key status
        api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if api_key:
            st.success("🔑 API Key configured")
        else:
            st.warning("🔑 No API Key - Using local generation")
    
    return True

# Custom CSS for modern styling
def apply_custom_styling():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&family=Nunito+Sans:wght@300;400;500;600&display=swap');
    
    /* Global Styles */
    .stApp {
        font-family: 'Nunito Sans', sans-serif;
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
        min-height: 100vh;
    }
    
    /* Main Container */
    .main .block-container {
        padding-top: 2rem;
        background: rgba(30, 30, 46, 0.95);
        border-radius: 20px;
        margin: 1rem auto;
        max-width: 1200px;
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.3);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        text-align: center;
    }
    
    /* Headings */
    h1, h2, h3, .stSelectbox label, .stCheckbox label {
        font-family: 'Poppins', sans-serif !important;
        font-weight: 600 !important;
        text-transform: uppercase !important;
        letter-spacing: 1px !important;
        color: #ffffff !important;
    }
    
    h1 {
        background: linear-gradient(135deg, #00d4ff, #5b86e5);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 2.5rem !important;
        text-align: center;
        margin-bottom: 1rem !important;
        text-shadow: 0 0 20px rgba(0, 212, 255, 0.3);
    }
    
    h2 {
        color: #ffffff !important;
        border-bottom: 2px solid #00d4ff;
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem !important;
        text-shadow: 0 0 10px rgba(0, 212, 255, 0.5);
        text-align: center;
    }
    
    h3 {
        text-align: center;
        color: #ffffff !important;
    }
    
    /* Sidebar Styling */
    .css-1d391kg {
        background: linear-gradient(180deg, #0f1419 0%, #1a1a2e 100%);
        border-radius: 0 20px 20px 0;
        border-right: 1px solid rgba(0, 212, 255, 0.3);
    }
    
    .sidebar .sidebar-content {
        background: transparent;
    }
    
    .sidebar h2, .sidebar h3, .sidebar label {
        color: #ffffff !important;
    }
    
    /* Input Fields with Dark Mode Glow */
    .stTextInput input, .stTextArea textarea, .stSelectbox select {
        border: 2px solid #2a2a3e !important;
        border-radius: 12px !important;
        padding: 12px !important;
        font-family: 'Nunito Sans', sans-serif !important;
        transition: all 0.3s ease !important;
        background: rgba(26, 26, 46, 0.8) !important;
        color: #ffffff !important;
    }
    
    .stTextInput input:focus, .stTextArea textarea:focus, .stSelectbox select:focus {
        border-color: #00d4ff !important;
        box-shadow: 0 0 20px rgba(0, 212, 255, 0.4) !important;
        transform: translateY(-2px) !important;
        background: rgba(26, 26, 46, 0.9) !important;
    }
    
    /* Buttons */
    .stButton button {
        background: linear-gradient(135deg, #00d4ff, #5b86e5) !important;
        color: white !important;
        border: none !important;
        border-radius: 25px !important;
        padding: 12px 30px !important;
        font-family: 'Poppins', sans-serif !important;
        font-weight: 600 !important;
        text-transform: uppercase !important;
        letter-spacing: 1px !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 8px 15px rgba(0, 212, 255, 0.3) !important;
    }
    
    .stButton button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 15px 25px rgba(0, 212, 255, 0.5) !important;
    }
    
    /* Primary Button */
    .stButton[data-baseweb="button"][kind="primary"] button {
        background: linear-gradient(135deg, #ff6b6b, #ee5a24) !important;
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0% { box-shadow: 0 8px 15px rgba(255, 107, 107, 0.3); }
        50% { box-shadow: 0 15px 25px rgba(255, 107, 107, 0.5); }
        100% { box-shadow: 0 8px 15px rgba(255, 107, 107, 0.3); }
    }
    
    /* Navigation buttons removed to fix React errors */
    
    /* Sticky Action Bar */
    .sticky-actions {
        position: sticky;
        top: 0;
        background: rgba(26, 26, 46, 0.95);
        backdrop-filter: blur(10px);
        padding: 15px;
        border-radius: 15px;
        margin-bottom: 20px;
        box-shadow: 0 5px 15px rgba(0, 0, 0, 0.3);
        border: 1px solid rgba(0, 212, 255, 0.2);
        z-index: 100;
    }
    
    /* Animated Question Cards */
    .question-card {
        background: linear-gradient(135deg, #1e1e2e, #252540);
        border-radius: 15px;
        padding: 20px;
        margin: 15px 0;
        border-left: 5px solid #00d4ff;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.4);
        transition: all 0.3s ease;
        animation: slideInUp 0.6s ease-out;
        border: 1px solid rgba(0, 212, 255, 0.1);
    }
    
    .question-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 35px rgba(0, 212, 255, 0.2);
        border-left-color: #ff6b6b;
    }
    
    @keyframes slideInUp {
        from {
            opacity: 0;
            transform: translateY(30px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* Expanders */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, #00d4ff, #5b86e5) !important;
        color: white !important;
        border-radius: 12px !important;
        font-family: 'Poppins', sans-serif !important;
        font-weight: 600 !important;
        text-transform: uppercase !important;
        letter-spacing: 1px !important;
    }
    
    .streamlit-expanderContent {
        background: #1e1e2e !important;
        border-radius: 0 0 12px 12px !important;
        border: 2px solid #00d4ff !important;
        border-top: none !important;
    }
    
    /* Success/Error Messages */
    .stSuccess {
        background: linear-gradient(135deg, #00d4aa, #01a3a4) !important;
        color: white !important;
        border-radius: 12px !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(0, 212, 170, 0.3) !important;
    }
    
    .stError {
        background: linear-gradient(135deg, #ff6b6b, #ee5a24) !important;
        color: white !important;
        border-radius: 12px !important;
        font-weight: 600 !important;
        box-shadow: 0 4px 15px rgba(255, 107, 107, 0.3) !important;
    }
    
    /* Progress Bars */
    .stProgress .st-bo {
        background: linear-gradient(135deg, #00d4ff, #5b86e5) !important;
        border-radius: 10px !important;
    }
    
    /* Sliders */
    .stSlider .st-bf {
        background: linear-gradient(135deg, #00d4ff, #5b86e5) !important;
    }
    
    /* Checkboxes */
    .stCheckbox input:checked + span {
        background: linear-gradient(135deg, #00d4ff, #5b86e5) !important;
    }
    
    /* Radio Buttons */
    .stRadio label {
        background: rgba(26, 26, 46, 0.8) !important;
        border-radius: 10px !important;
        padding: 10px !important;
        margin: 5px 0 !important;
        transition: all 0.3s ease !important;
        color: #ffffff !important;
        border: 1px solid rgba(0, 212, 255, 0.2) !important;
    }
    
    .stRadio label:hover {
        background: rgba(0, 212, 255, 0.1) !important;
        transform: translateX(5px) !important;
        border-color: #00d4ff !important;
    }
    
    /* File Uploader */
    .stFileUploader {
        background: linear-gradient(135deg, #1e1e2e, #252540) !important;
        border: 2px dashed #00d4ff !important;
        border-radius: 15px !important;
        padding: 30px !important;
        text-align: center !important;
        transition: all 0.3s ease !important;
        color: #ffffff !important;
        margin: 0 auto !important;
        max-width: 600px !important;
    }
    
    /* Center align all content */
    .stApp > div > div > div > div {
        text-align: center !important;
    }
    
    /* Center form elements */
    .stSelectbox, .stSlider, .stTextInput, .stTextArea, .stNumberInput {
        text-align: center !important;
        margin: 0 auto !important;
        max-width: 600px !important;
    }
    
    /* Center buttons */
    .stButton {
        text-align: center !important;
        margin: 1rem auto !important;
        display: block !important;
    }
    
    /* Center columns */
    .row-widget {
        justify-content: center !important;
    }
    
    /* Center labels */
    label {
        text-align: center !important;
        display: block !important;
        width: 100% !important;
    }
    
    /* Center content sections */
    .element-container {
        text-align: center !important;
    }
    
    .stFileUploader:hover {
        border-color: #ff6b6b !important;
        background: rgba(0, 212, 255, 0.1) !important;
        box-shadow: 0 0 20px rgba(0, 212, 255, 0.2) !important;
    }
    
    /* Download Buttons */
    .download-section {
        background: linear-gradient(135deg, #1e1e2e, #252540);
        border-radius: 15px;
        padding: 25px;
        margin: 20px 0;
        border: 2px solid rgba(0, 212, 255, 0.3);
        text-align: center;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.3);
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .main .block-container {
            margin: 0.5rem;
            border-radius: 15px;
        }
        
        /* Responsive adjustments for removed floating buttons */
    }
    </style>
    """, unsafe_allow_html=True)

def extract_text_from_pdf(pdf_file):
    """
    Extract clean, readable text from uploaded PDF file using pdfplumber
    """
    if pdfplumber is None:
        raise Exception("pdfplumber is required for PDF processing. Please install it.")
    
    try:
        text_content = []
        
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        extracted_text = '\n\n'.join(text_content)
        
        if not extracted_text.strip():
            raise Exception("No text could be extracted from the PDF. The PDF might contain only images or be password protected.")
        
        return extracted_text
    
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_article(article_text):
    """
    Process and clean article text for quiz generation
    """
    if not article_text or not article_text.strip():
        raise Exception("Article text cannot be empty.")
    
    # Basic text cleaning
    cleaned_text = article_text.strip()
    
    # Remove excessive whitespace and normalize line breaks
    cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
    cleaned_text = re.sub(r' +', ' ', cleaned_text)
    
    if len(cleaned_text) < 100:
        raise Exception("Article text is too short. Please provide at least 100 characters for meaningful quiz generation.")
    
    return cleaned_text

def generate_questions_with_gemini_ai(text_content, num_mc=5, num_tf=5, difficulty="medium"):
    """
    Generate quiz questions using Gemini AI
    """
    try:
        from gemini_quiz_generator import GeminiQuizGenerator
        
        generator = GeminiQuizGenerator()
        quiz_data = generator.generate_quiz(text_content, difficulty, num_mc, num_tf)
        
        return quiz_data
        
    except Exception as e:
        # Fallback to local generation if Gemini fails
        error_msg = str(e)
        if "GEMINI_API_KEY" in error_msg or "GOOGLE_API_KEY" in error_msg:
            st.info("🔑 To use Gemini AI, set your GEMINI_API_KEY or GOOGLE_API_KEY environment variable. Using local generation for now.")
        else:
            st.warning(f"Gemini AI temporarily unavailable: {error_msg}. Using local generation as backup.")
        return create_intelligent_questions(text_content, num_mc, num_tf, difficulty, ["multiple_choice", "true_false"])

def create_intelligent_questions(content, num_mc=5, num_tf=5, difficulty="medium", question_types=["multiple_choice", "true_false"]):
    """
    Create intelligent quiz questions using local content analysis
    """
    # Process content into sentences and extract key information
    sentences = [s.strip() for s in re.split(r'[.!?]+', content) if len(s.strip()) > 20]
    words = content.split()
    paragraphs = [p.strip() for p in content.split('\n\n') if len(p.strip()) > 50]
    
    # Extract key terms, numbers, and facts
    common_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'who', 'boy', 'did', 'she', 'use', 'way', 'when', 'what', 'with', 'this', 'that', 'have', 'from', 'they', 'know', 'want', 'been', 'good', 'much', 'some', 'time', 'very', 'were', 'will', 'your', 'about', 'after', 'before', 'other', 'right', 'their', 'there', 'these', 'which', 'would', 'could', 'should', 'often', 'called'}
    
    # Extract key terms (proper nouns, important concepts)
    key_terms = []
    for word in words:
        clean_word = re.sub(r'[^\w]', '', word)
        if (len(clean_word) > 4 and 
            clean_word.lower() not in common_words and
            (clean_word[0].isupper() or len(clean_word) > 7)):
            key_terms.append(clean_word)
    
    # Remove duplicates and take top terms
    key_terms = list(dict.fromkeys(key_terms))[:15]
    
    # Generate Multiple Choice Questions
    multiple_choice = []
    
    # Only generate multiple choice if requested
    if "multiple_choice" in question_types and key_terms:
        for i in range(min(num_mc, len(key_terms))):
            term = key_terms[i]
            
            # Generate question based on difficulty
            if difficulty == "easy":
                question = f"What is {term}?"
                correct_option = "A key concept discussed in the text"
                wrong_options = [
                    "Something not mentioned",
                    "An unrelated topic", 
                    "A minor detail"
                ]
            elif difficulty == "medium":
                question = f"According to the content, what is the significance of {term}?"
                correct_option = "It plays a crucial role in the topic discussed"
                wrong_options = [
                    "It has minimal importance",
                    "It contradicts the main ideas",
                    "It is only mentioned in passing"
                ]
            else:  # hard
                question = f"Analyze the role of {term} within the broader context of this topic?"
                correct_option = "It represents a critical factor that influences multiple aspects of the subject"
                wrong_options = [
                    "It serves as a peripheral element with limited interconnections",
                    "It contradicts the established theoretical framework", 
                    "It represents an anomaly that challenges conventional understanding"
                ]
            
            # Shuffle options
            all_options = [correct_option] + wrong_options
            random.shuffle(all_options)
            correct_index = all_options.index(correct_option)
            correct_letter = ['A', 'B', 'C', 'D'][correct_index]
            
            formatted_options = [f"{['A', 'B', 'C', 'D'][j]}) {opt}" for j, opt in enumerate(all_options)]
            
            multiple_choice.append({
                "question": question,
                "options": formatted_options,
                "correct_answer": correct_letter,
                "explanation": f"Based on the content, {term} is discussed as {correct_option.lower()}."
            })
    
    # Generate True/False Questions
    true_false = []
    
    # Only generate true/false if requested
    if "true_false" in question_types and sentences:
        for i in range(min(num_tf, len(sentences))):
            sentence = sentences[i]
            
            # Adjust complexity based on difficulty
            if difficulty == "easy":
                statement = sentence[:100] + "..." if len(sentence) > 100 else sentence
                true_false.append({
                    "question": statement,
                    "correct_answer": "True",
                    "explanation": "This statement is directly from the content."
                })
            elif difficulty == "medium":
                statement = sentence[:150] + "..." if len(sentence) > 150 else sentence
                true_false.append({
                    "question": statement,
                    "correct_answer": "True", 
                    "explanation": "This statement is directly supported by the provided content."
                })
            else:  # hard
                if key_terms:
                    term = random.choice(key_terms)
                    statement = f"The text implies that {term} functions as a significant component within the contextual framework presented"
                    true_false.append({
                        "question": statement,
                        "correct_answer": "True",
                        "explanation": f"While not explicitly stated in these exact words, the content supports this interpretation regarding {term}."
                    })
    
    # Return only requested question types
    result = {}
    if "multiple_choice" in question_types:
        result["multiple_choice"] = multiple_choice[:num_mc]
    if "true_false" in question_types:
        result["true_false"] = true_false[:num_tf]
    
    return result

def render_quiz_form(questions_data):
    """
    Render the quiz questions in an editable form with enhanced styling
    """
    st.header("📝 Generated Quiz Questions")
    
    if 'edited_questions' not in st.session_state:
        st.session_state.edited_questions = questions_data.copy()
    
    # Multiple Choice Questions
    mc_questions = st.session_state.edited_questions.get('multiple_choice', [])
    if mc_questions:
        st.subheader("Multiple Choice Questions")
        
        for i, q in enumerate(mc_questions):
            # Animated question card
            st.markdown(f'<div class="question-card">', unsafe_allow_html=True)
            with st.expander(f"Question {i+1}: {q['question'][:50]}..."):
                # Question text
                new_question = st.text_area(f"Question {i+1}", value=q['question'], key=f"mc_q_{i}")
                
                # Options
                st.write("Options:")
                new_options = []
                for j, option in enumerate(q['options']):
                    new_option = st.text_input(f"Option {j+1}", value=option, key=f"mc_opt_{i}_{j}")
                    new_options.append(new_option)
                
                # Correct answer
                correct_choices = ['A', 'B', 'C', 'D']
                current_correct = q['correct_answer']
                correct_index = correct_choices.index(current_correct) if current_correct in correct_choices else 0
                new_correct = st.selectbox(f"Correct Answer", correct_choices, index=correct_index, key=f"mc_ans_{i}")
                
                # Explanation
                new_explanation = st.text_area(f"Explanation", value=q['explanation'], key=f"mc_exp_{i}")
                
                # Update question in session state
                st.session_state.edited_questions['multiple_choice'][i] = {
                    'question': new_question,
                    'options': new_options,
                    'correct_answer': new_correct,
                    'explanation': new_explanation
                }
                
                # Delete button
                if st.button(f"Delete Question {i+1}", key=f"del_mc_{i}"):
                    del st.session_state.edited_questions['multiple_choice'][i]
                    st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # True/False Questions
    tf_questions = st.session_state.edited_questions.get('true_false', [])
    if tf_questions:
        st.subheader("True/False Questions")
        
        for i, q in enumerate(tf_questions):
            # Animated question card
            st.markdown(f'<div class="question-card">', unsafe_allow_html=True)
            with st.expander(f"T/F Question {i+1}: {q['question'][:50]}..."):
                # Question text
                new_question = st.text_area(f"T/F Question {i+1}", value=q['question'], key=f"tf_q_{i}")
                
                # Correct answer
                current_correct = q['correct_answer']
                tf_choices = ['True', 'False']
                correct_index = tf_choices.index(current_correct) if current_correct in tf_choices else 0
                new_correct = st.selectbox(f"Correct Answer", tf_choices, index=correct_index, key=f"tf_ans_{i}")
                
                # Explanation
                new_explanation = st.text_area(f"Explanation", value=q['explanation'], key=f"tf_exp_{i}")
                
                # Update question in session state
                st.session_state.edited_questions['true_false'][i] = {
                    'question': new_question,
                    'correct_answer': new_correct,
                    'explanation': new_explanation
                }
                
                # Delete button
                if st.button(f"Delete T/F Question {i+1}", key=f"del_tf_{i}"):
                    del st.session_state.edited_questions['true_false'][i]
                    st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)

def export_quiz(questions_data, format_type="txt"):
    """
    Export quiz questions to downloadable format
    """
    if format_type == "txt":
        return export_quiz_txt(questions_data)
    elif format_type == "docx":
        return export_quiz_docx(questions_data)

def export_quiz_txt(questions_data):
    """
    Export quiz to TXT format
    """
    content = "QUIZ QUESTIONS\n" + "="*50 + "\n\n"
    
    # Multiple Choice Questions
    mc_questions = questions_data.get('multiple_choice', [])
    if mc_questions:
        content += "MULTIPLE CHOICE QUESTIONS:\n" + "-"*30 + "\n\n"
        for i, q in enumerate(mc_questions):
            content += f"{i+1}. {q['question']}\n"
            for option in q['options']:
                content += f"   {option}\n"
            content += f"   Correct Answer: {q['correct_answer']}\n"
            content += f"   Explanation: {q['explanation']}\n\n"
    
    # True/False Questions
    tf_questions = questions_data.get('true_false', [])
    if tf_questions:
        content += "TRUE/FALSE QUESTIONS:\n" + "-"*30 + "\n\n"
        for i, q in enumerate(tf_questions):
            content += f"{len(mc_questions) + i + 1}. {q['question']}\n"
            content += f"   Correct Answer: {q['correct_answer']}\n"
            content += f"   Explanation: {q['explanation']}\n\n"
    
    return content.encode('utf-8')

def export_quiz_docx(questions_data):
    """
    Export quiz to DOCX format
    """
    if Document is None:
        raise Exception("python-docx is required for DOCX export. Please install it.")
    
    doc = Document()
    doc.add_heading('Quiz Questions', 0)
    
    # Multiple Choice Questions
    mc_questions = questions_data.get('multiple_choice', [])
    if mc_questions:
        doc.add_heading('Multiple Choice Questions', level=1)
        for i, q in enumerate(mc_questions):
            doc.add_paragraph(f"{i+1}. {q['question']}")
            for option in q['options']:
                doc.add_paragraph(f"    {option}", style='List Bullet')
            doc.add_paragraph(f"Correct Answer: {q['correct_answer']}", style='Normal')
            doc.add_paragraph(f"Explanation: {q['explanation']}", style='Normal')
            doc.add_paragraph()  # Empty line
    
    # True/False Questions
    tf_questions = questions_data.get('true_false', [])
    if tf_questions:
        doc.add_heading('True/False Questions', level=1)
        for i, q in enumerate(tf_questions):
            doc.add_paragraph(f"{len(mc_questions) + i + 1}. {q['question']}")
            doc.add_paragraph(f"Correct Answer: {q['correct_answer']}")
            doc.add_paragraph(f"Explanation: {q['explanation']}")
            doc.add_paragraph()  # Empty line
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()

# Main App
def main():
    # Check platform compatibility and show status
    check_platform_compatibility()
    
    # Apply custom styling
    apply_custom_styling()
    
    # Navigation helper (removed floating buttons to fix React errors)
    
    st.title("📚 AI Quiz Generator")
    st.markdown("Upload PDF/CSV files or paste article text to generate customizable quiz questions with Google's advanced Gemini AI")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        st.subheader("Gemini AI Generation")
        st.success("✅ Gemini AI enabled")
        st.info("Powered by Google's advanced Gemini AI for high-quality question generation")
        
        st.subheader("Supported Content Types")
        st.write("📄 **PDF Files**: Chapters, research papers, textbooks")
        st.write("📊 **CSV Data**: Datasets, tables, statistical information")
        st.write("📰 **Articles**: Blog posts, news articles, web content")
        st.write("📖 **Text Content**: Any educational or informational text")
        
        st.subheader("Quiz Settings")
        
        # Difficulty selection
        difficulty = st.radio(
            "Difficulty Level:",
            ["easy", "medium", "hard"],
            index=1,
            help="Easy: Simple vocabulary and basic questions\nMedium: Standard complexity\nHard: Advanced vocabulary and complex analysis"
        )
        
        # Question type selection
        st.write("**Question Types:**")
        include_mc = st.checkbox("Multiple Choice Questions", value=True)
        include_tf = st.checkbox("True/False Questions", value=True)
        
        # Question numbers (only show if type is selected)
        if include_mc:
            num_mc = st.slider("Number of Multiple Choice Questions", 1, 15, 5)
        else:
            num_mc = 0
            
        if include_tf:
            num_tf = st.slider("Number of True/False Questions", 1, 15, 5)
        else:
            num_tf = 0
        
        # Validation
        if not include_mc and not include_tf:
            st.error("Please select at least one question type!")
            st.stop()
    
    # Input method selection
    st.header("📝 Choose Input Method")
    st.markdown("Select how you want to provide content:")
    
    # Create three equal columns for horizontal layout
    col1, col2, col3 = st.columns(3)
    
    # Get current selection or default
    current_method = st.session_state.get('input_method', "📄 Upload PDF File")
    
    with col1:
        pdf_type = "primary" if current_method == "📄 Upload PDF File" else "secondary"
        pdf_selected = st.button("📄 Upload PDF File", use_container_width=True, type=pdf_type)
    
    with col2:
        csv_type = "primary" if current_method == "📊 Upload CSV Data" else "secondary"
        csv_selected = st.button("📊 Upload CSV Data", use_container_width=True, type=csv_type)
    
    with col3:
        article_type = "primary" if current_method == "📰 Paste Article Text" else "secondary"
        article_selected = st.button("📰 Paste Article Text", use_container_width=True, type=article_type)
    
    # Store selection in session state
    if pdf_selected:
        st.session_state.input_method = "📄 Upload PDF File"
    elif csv_selected:
        st.session_state.input_method = "📊 Upload CSV Data"
    elif article_selected:
        st.session_state.input_method = "📰 Paste Article Text"
    
    # Update current selection
    input_method = st.session_state.get('input_method', "📄 Upload PDF File")
    
    extracted_text = None
    
    if input_method == "📄 Upload PDF File":
        st.subheader("PDF File Upload")
        uploaded_file = st.file_uploader(
            "Choose a PDF file:",
            type=['pdf'],
            help="Upload a PDF chapter containing text content for quiz generation"
        )
        
        if uploaded_file is not None:
            try:
                # Extract text
                with st.spinner("Extracting text from PDF..."):
                    extracted_text = extract_text_from_pdf(uploaded_file)
                
                st.success(f"✅ Text extracted successfully! ({len(extracted_text)} characters)")
                
            except Exception as e:
                st.error(f"Error processing PDF: {str(e)}")
    
    elif input_method == "📊 Upload CSV Data":
        st.subheader("CSV Data Upload")
        uploaded_file = st.file_uploader(
            "Choose a CSV file:",
            type=['csv'],
            help="Upload a CSV file containing data for quiz generation"
        )
        
        if uploaded_file is not None:
            try:
                # Process CSV file
                from file_processor import FileProcessor
                processor = FileProcessor()
                
                with st.spinner("Processing CSV data..."):
                    extracted_text = processor.process_file(uploaded_file)
                
                st.success(f"✅ CSV data processed successfully! ({len(extracted_text)} characters)")
                
            except Exception as e:
                st.error(f"Error processing CSV: {str(e)}")
    
    elif input_method == "📰 Paste Article Text":
        st.subheader("Article Text Input")
        article_text = st.text_area(
            "Paste your article content here:",
            height=300,
            placeholder="Paste the article text, blog post, or any written content you want to create quiz questions from...",
            help="You can paste articles from websites, documents, or any text content"
        )
        
        if article_text:
            try:
                extracted_text = extract_text_from_article(article_text)
                st.success(f"✅ Article processed successfully! ({len(extracted_text)} characters)")
                
            except Exception as e:
                st.error(f"Error processing article: {str(e)}")
    
    # Show preview and generate questions if text is available
    if extracted_text:
        # Show preview
        with st.expander("📖 Preview Content"):
            st.text_area("Content:", extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text, height=200, disabled=True)
        
        # Generate questions
        if st.button("🎯 Generate Quiz Questions", type="primary"):
            try:
                # Determine question types to generate
                question_types = []
                if include_mc:
                    question_types.append("multiple_choice")
                if include_tf:
                    question_types.append("true_false")
                
                with st.spinner(f"Analyzing content with Gemini AI and generating {difficulty} level quiz questions..."):
                    questions_data = generate_questions_with_gemini_ai(
                        extracted_text, 
                        num_mc, 
                        num_tf, 
                        difficulty
                    )
                
                st.success("✅ Quiz questions generated successfully!")
                
                # Store in session state
                st.session_state.questions_generated = True
                st.session_state.original_questions = questions_data
                st.session_state.edited_questions = questions_data.copy()
                
            except Exception as e:
                st.error(f"Error generating questions: {str(e)}")
    
    # Show questions if generated
    if st.session_state.get('questions_generated', False):
        # Sticky Action Bar
        st.markdown("""
        <div class="sticky-actions">
            <div style="display: flex; justify-content: space-between; align-items: center;">
                <div style="font-family: 'Poppins', sans-serif; font-weight: 600; color: #ffffff;">
                    Quiz Generated ✅ | Ready for Review & Export
                </div>
                <div style="display: flex; gap: 10px;">
                    <button onclick="window.scrollTo(0, document.querySelector('[data-testid=\\"stHeader\\"]').offsetTop)" 
                            style="padding: 8px 16px; background: linear-gradient(135deg, #00d4ff, #5b86e5); color: white; border: none; border-radius: 20px; cursor: pointer; box-shadow: 0 4px 15px rgba(0, 212, 255, 0.3);">
                        📝 Edit Questions
                    </button>
                    <button onclick="window.scrollTo(0, document.body.scrollHeight)" 
                            style="padding: 8px 16px; background: linear-gradient(135deg, #ff6b6b, #ee5a24); color: white; border: none; border-radius: 20px; cursor: pointer; box-shadow: 0 4px 15px rgba(255, 107, 107, 0.3);">
                        📥 Export Quiz
                    </button>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        render_quiz_form(st.session_state.edited_questions)
        
        # Export options with enhanced styling
        st.markdown('<div class="download-section">', unsafe_allow_html=True)
        st.header("📥 Export Quiz")
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Download as TXT"):
                txt_data = export_quiz(st.session_state.edited_questions, "txt")
                if txt_data:
                    st.download_button(
                        label="Download TXT File",
                        data=txt_data,
                        file_name="quiz.txt",
                        mime="text/plain"
                    )
                else:
                    st.error("Failed to generate TXT file")
        
        with col2:
            if st.button("📄 Download as DOCX"):
                try:
                    docx_data = export_quiz(st.session_state.edited_questions, "docx")
                    if docx_data:
                        st.download_button(
                            label="Download DOCX File", 
                            data=docx_data,
                            file_name="quiz.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("Failed to generate DOCX file")
                except Exception as e:
                    st.error(f"Error creating DOCX: {str(e)}")
        
        st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
